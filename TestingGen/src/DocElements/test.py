from typing import Optional

from src.element_wizard import ElementWizard
from src.path_tree import PathTree
from .doc_element import IDocElement

from docx.shared import Pt
from docx.shared import RGBColor
from docx.document import Document as DocumentType

from src.cancelable_input import CancelableInput
from src.context import Context

class Test(IDocElement):
	allowed_elements = ["Paragraph", "Picture", "Screenshot"]

	def __init__(self, title: str, passed: bool, description: Optional[str] = None, expected_output: Optional[str] = None, prefix_elements: Optional[list[IDocElement]] = None, suffix_elements: Optional[list[IDocElement]] = None):
		super().__init__()

		self.title = title
		self.passed = passed
		self.description = description
		self.expected_output = expected_output
		self.prefix_elements = prefix_elements
		self.suffix_elements = suffix_elements

		super().save_document()

	def serialize(self) -> dict:
		data = {
			"type": self.get_type(),
			"title": self.title,
			"passed": self.passed,
		}

		if self.description:
			data["description"] = self.description

		if self.expected_output:
			data["expected_output"] = self.expected_output

		if self.prefix_elements:
			data["prefix_elements"] = [e.serialize() for e in self.prefix_elements]

		if self.suffix_elements:
			data["suffix_elements"] = [e.serialize() for e in self.suffix_elements]

		return data

	@classmethod
	def deserialize(cls, data) -> 'Test':
		title = data["title"]
		passed = data["passed"]

		description = data["description"] if "description" in data else None
		expected_output = data["expected_output"] if "expected_output" in data else None

		prefix_elements = [IDocElement.deserialize(e) for e in data["prefix_elements"]] if "prefix_elements" in data else None
		suffix_elements = [IDocElement.deserialize(e) for e in data["suffix_elements"]] if "suffix_elements" in data else None

		return cls(title, passed, description, expected_output, prefix_elements, suffix_elements)
	
	def doc_gen(self, doc: DocumentType, test_num: Optional[int] = None, subtest_num: Optional[int] = None):
		title = self.title
		if test_num:
			prefix = f"Test {test_num}"

			if subtest_num:
				prefix += f".{subtest_num}: "
			else:
				prefix += ": "

			title = prefix + title

		title_run = doc.add_paragraph().add_run(title)
		title_run.font.size = Pt(14)
		title_run.underline = True

		if self.description:
			doc.add_paragraph(self.description)

		if self.expected_output:
			expected_paragraph = doc.add_paragraph()
			expected_paragraph.add_run("Expected Output:").underline = True
			expected_paragraph.add_run(f" {self.expected_output}")

		if self.prefix_elements:
			for e in self.prefix_elements:
				e.doc_gen(doc)

		passed_col = RGBColor(0, 128, 0) if self.passed else RGBColor(255, 0, 0)

		passed_paragraph = doc.add_paragraph()
		passed_run_pre = passed_paragraph.add_run("Result: ")
		passed_run_pre.font.color.rgb = passed_col
		passed_run_pre.font.bold = True

		passed_run = passed_paragraph.add_run("PASSED" if self.passed else "FAILED")
		passed_run.font.color.rgb = passed_col
		passed_run.font.bold = True
		passed_run.font.underline = True

		if self.suffix_elements:
			for e in self.suffix_elements:
				e.doc_gen(doc)

	@classmethod
	def wizard(cls) -> Optional['Test']:
		with PathTree("Test"):
			PathTree.cls()

			answers = CancelableInput.input_chain(
				"Test Title: ",
				"Description: ",
				"Expected Output: ",
				"Add prefix element? (y/n): ",
				"Add suffix element? (y/n): ",
			)

			if not answers: return None

			title = answers[0]

			description = answers[1].strip()
			description = description if description != "" else None

			expected_output = answers[2].strip()
			expected_output = expected_output if expected_output != "" else None

			prefix_elements = [] if answers[3].strip().lower() == "y" else None
			suffix_elements = [] if answers[4].strip().lower() == "y" else None

			instance = cls(title, passed, description, expected_output, prefix_elements, suffix_elements)

			if prefix_elements != None:
				ElementWizard.add_wizard(prefix_elements, cls.allowed_elements, status="Please select a prefix element to add:", cancel_option="Back", path_name="Add Prefix")
				instance.prefix_elements = prefix_elements
				instance.save_document()

			if suffix_elements != None:
				ElementWizard.add_wizard(suffix_elements, cls.allowed_elements, status="Please select a suffix element to add:", cancel_option="Back", path_name="Add Suffix")
				instance.suffix_elements = suffix_elements
				instance.save_document()

			PathTree.cls()
			passed = CancelableInput.input("Passed (y/n): ").strip().lower() == "y"
			instance.passed = passed

			return instance
	
	def edit(self):
		with PathTree("Test"):
			PathTree.cls()

			status = "Please select an option to edit:"

			while True:	
				options = ["Title", "Description", "Expected Output", "Passed"]

				prefix_option = "Prefix Elements"
				if self.prefix_elements:
					prefix_option += f" {[e for e in self.prefix_elements]}"

				suffix_option = "Suffix Elements"
				if self.suffix_elements:
					suffix_option += f" {[e for e in self.suffix_elements]}"

				options.append(prefix_option)
				options.append(suffix_option)

				option = ElementWizard.selection_wizard(options, status, "Back")

				PathTree.cls()

				match option:
					case -1:
						return
					case 0:
						out = CancelableInput.input("Title: ", self.title)

						if out:
							self.title = out
							status = "Title changed"
						else:
							status = "Title not changed"
					case 1:
						out = CancelableInput.input("Description: ", self.description)

						if out:
							self.description = out
							status = "Description changed"
						else:
							status = "Description not changed"
					case 2:
						out = CancelableInput.input("Expected Output: ", self.expected_output)

						if out:
							self.expected_output = out
							status = "Expected Output changed"
						else:
							status = "Expected Output not changed"
					case 3:
						out = CancelableInput.input("Passed (y/n): ", "y" if self.passed else "n")

						if out:
							self.passed = out.strip().lower() == "y"
							status = "Passed changed"
						else:
							status = "Passed not changed"
					case 4:
						self.prefix_elements = self.prefix_elements if self.prefix_elements else []
						ElementWizard.wizard(self.prefix_elements, "Prefix", allowed_types_to_add=self.__class__.allowed_elements)
					case 5:
						self.suffix_elements = self.suffix_elements if self.suffix_elements else []
						ElementWizard.wizard(self.suffix_elements, "Suffix", allowed_types_to_add=self.__class__.allowed_elements)
				

	def __str__(self) -> str:
		return f"Test ({self.title})"
	
	def __repr__(self) -> str:
		string = f"Test({repr(self.title)}, {repr(self.passed)}"

		if self.description: string += f", description={repr(self.description)}"
		if self.expected_output: string += f", expected_output={repr(self.expected_output)}"
		if self.prefix_elements: string += f", prefix_elements={repr(self.prefix_elements)}"
		if self.suffix_elements: string += f", suffix_elements={repr(self.suffix_elements)}"

		return string
		