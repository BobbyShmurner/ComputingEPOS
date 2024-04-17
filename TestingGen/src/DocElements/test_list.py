from typing import Optional

from src.DocElements.test import Test
from src.cancelable_input import CancelableInput
from src.element_wizard import ElementWizard
from src.path_tree import PathTree
from .doc_element import IDocElement

from docx.document import Document as DocumentType
from docx.shared import Pt

class TestList(IDocElement):
	test_list_count = 0

	def __init__(self, title: str, description: str, tests: list[Test]):
		super().__init__()

		self.title = title
		self.description = description
		self.tests = tests

		super().save_document()
		

	def serialize(self) -> dict:
		return {
			"type": self.get_type(),
			"title": self.title,
			"description": self.description,
			"tests": [t.serialize() for t in self.tests]
		}

	@classmethod
	def deserialize(cls, data) -> 'TestList':
		title = data["title"]
		description = data["description"]
		tests = [Test.deserialize(t) for t in data["tests"]]

		return cls(title, description, tests)
	
	def doc_gen(self, doc: DocumentType):
		cls = self.__class__
		cls.test_list_count += 1

		title_run = doc.add_paragraph().add_run(f"Test {cls.test_list_count}: {self.title}")
		title_run.underline = True
		title_run.font.size = Pt(16)
		title_run.bold = True

		doc.add_paragraph(self.description)

		for i, test in enumerate(self.tests):
			test.doc_gen(doc, cls.test_list_count, i + 1)
		

	@classmethod
	def wizard(cls) -> Optional['TestList']:
		with PathTree("Test List"):
			PathTree.cls()

			answers = CancelableInput.input_chain([
				"Test List Title: ",
				"Description: ",
			])

			if not answers: return None

			title = answers[0]
			description = answers[1].strip()
			tests = []

			instance = cls(title, description, tests)

			ElementWizard.add_wizard(tests, ["Test"], status="Add a test?", cancel_option="Back")

			instance.tests = tests
			instance.save_document()
			
			return instance
	
	def edit(self):
		with PathTree(str(self)):
			PathTree.cls()
			status = "Please select an option to edit:"

			while True:	
				options = ["Title", "Description", "Tests"]
				option = ElementWizard.selection_wizard(options, status, "Back")

				PathTree.cls()

				match option:
					case -1:
						return
					case 0:
						out = CancelableInput.input("Title: ", self.title)

						if out:
							self.title = out
							status = "Title changed"
						else:
							status = "Title not changed"
					case 1:
						out = CancelableInput.input("Description: ", self.description)

						if out:
							self.description = out
							status = "Description changed"
						else:
							status = "Description not changed"
					case 2:
						ElementWizard.wizard(self.tests, "Tests", allowed_types_to_add=["Test"])

	@classmethod
	def display_name(cls) -> str:
		return "Test List"

	def __str__(self) -> str:
		return f"Test List ({self.title})"
	
	def __repr__(self) -> str:
		return f"TestList({repr(self.title)}, {repr(self.description)}, {repr(self.tests)})"