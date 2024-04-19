import msvcrt
import os
import shutil
import PIL
import cv2
import time
import docx
import win32con
import win32gui
import numpy as np
import win32com.client
from io import BytesIO
from typing import Optional
from PIL import Image, ImageGrab
from docx.document import Document as DocumentType
from docx.text.run import Run

from typing import TYPE_CHECKING

if TYPE_CHECKING:
	from docx.oxml.xmlchemy import BaseOxmlElement

from src.cancelable_input import CancelableInput
from src.element_wizard import ElementWizard
from src.path_tree import PathTree

from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .doc_element import IDocElement
from ..context import Context

class Widget(IDocElement):
	def __init__(self, name: str):
		super().__init__()

		self.name = name

		super().save_document()

	def serialize(self) -> dict:
		data = {
			"type": self.get_type(),
			"name": self.name,
			
		}

		return data

	@classmethod
	def deserialize(cls, data: dict) -> 'Widget':
		name = data["name"]

		return cls(name)
	
	def doc_gen(self, doc: DocumentType, run: Optional[Run] = None):
		try:
			widget_doc: DocumentType = docx.Document(self.widget_path())

			with open(self.widget_path() + ".xml", 'w+') as file:
				file.write(widget_doc.element.body.xml)


			if run == None:
				current_para = doc.paragraphs[-1] if len(doc.paragraphs) > 0 else doc.add_paragraph()
				run = current_para.runs[-1] if len(current_para.runs) > 0 else current_para.add_run()
			else:
				current_para = Paragraph(run.element.getparent(), doc)
			prev_para = current_para.insert_paragraph_before()

			for r in current_para.runs:
				if r.element == run.element:
					break

				prev_para._element.append(r._element)

			for element in widget_doc.element.body:
				element: BaseOxmlElement = element

				current_para._element.insert_element_before(element)
		except Exception as e:
			print("Error when generating document widget!\n\n")
			input(e)

	def widget_path(self) -> str:
		return os.path.join(self.widget_dir(), f"{self.name}.docx")
	
	@classmethod
	def widget_dir(cls) -> str:
		return Context._instance.get_content_path(f"Widgets/")

	@classmethod
	def wizard(cls) -> Optional['Widget']:
		with PathTree("Widget"):
			PathTree.cls()

			status = "Use existing Widget?"
			choices = ["Existing", "Create New"]

			widget_name: str = None

			while widget_name == None:
				choice = ElementWizard.selection_wizard(choices, status)
				match choice:
					case -1:
						return None
					case 0:
						
						widgets = [".".join(widget.split(".")[:-1]) for widget in os.listdir(cls.widget_dir())]
						status = "Please Select a Widget"

						
						while widget_name == None:
							choice = ElementWizard.selection_wizard(widgets, status)

							match choice:
								case -1:
									break
								case _:
									widget_name = widgets[choice]
					case 1:
						with PathTree("New"):
							answers = CancelableInput.input_chain("Widget Name: ", "Widget Path: ")
							if not answers: break

							widget_name = answers[0]
							widget_path = answers[1]

							if not os.path.exists(widget_path):
								status = "Path does not exist"
								break

							os.makedirs(cls.widget_dir(), exist_ok=True)
							shutil.copy(widget_path, os.path.join(cls.widget_dir(), f"{widget_name}.docx"))

		return cls(widget_name)

	def edit(self):
		with PathTree(str(self)):
			PathTree.cls()
			

	def __str__(self) -> str:
		return f"Widget (Name: {self.name})"
	
	def __repr__(self) -> str:
		return f"Widget({repr(self.name)})"