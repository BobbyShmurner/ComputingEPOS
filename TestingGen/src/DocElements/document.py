import os
import json
import time
import msvcrt
from pathlib import Path
from typing import Optional

from src.context import Context
from src.element_wizard import ElementWizard
from src.path_tree import PathTree
from .doc_element import IDocElement
from src.DocElements import Picture, Paragraph

import docx
from docx.shared import Pt, Mm
from docx.document import Document as DocumentType

class Document(IDocElement):
	def __init__(self, elements: list[IDocElement] = [], font: Optional[str] = None, font_size: Optional[int] = None):
		super().__init__()

		self.elements = elements
		self.font = font
		self.font_size = font_size

	def add_element(self, element: 'IDocElement'):
		self.elements.append(element)

	def serialize(self) -> dict:
		data = {
			"type": self.get_type(),
		}

		if self.font:
			data["font"] = self.font

		if self.font_size:
			data["font_size"] = self.font_size

		data["elements"] = [e.serialize() for e in self.elements]

		return data

	@classmethod
	def deserialize(cls, data) -> 'Document':
		elements_data = data["elements"]
		elements = []

		font = data["font"] if "font" in data else None
		font_size = data["font_size"] if "font_size" in data else None
		
		for data in elements_data:
			elements.append(IDocElement.deserialize(data))

		return cls(elements, font, font_size)
	
	def serialize_to_disk(self):
		while True:
			try:
				path = self.context.get_content_path("Document.json")
				os.makedirs(os.path.dirname(path), exist_ok=True)

				back = ""
				if os.path.exists(path):
					with open(path, "r+") as file:
						back = file.read()

				try:
					with open(path, "w+") as file:
						file.write(json.dumps(self.serialize(), indent=4))
				except:
					with open(path, "w+") as file:
						file.write(back)

					raise

				break
			except Exception as e:
				print(e)
				print("\nException occured during saving, retrying...")

	@classmethod
	def deserialize_from_disk(cls) -> 'Document':
		path = Context._instance.get_content_path("Document.json")

		with open(path, "r") as file:
			data = json.loads(file.read())

		return cls.deserialize(data)
	
	def doc_gen(self, path: str = "out.docx"):
		doc: DocumentType = docx.Document()	
		
		# A4 Page Setup (Copied from Word)
		section = doc.sections[0]
		section.page_height = Mm(297)
		section.page_width = Mm(210)
		section.left_margin = Mm(25.4)
		section.right_margin = Mm(25.4)
		section.top_margin = Mm(25.4)
		section.bottom_margin = Mm(25.4)
		section.header_distance = Mm(12.5)
		section.footer_distance = Mm(12.5)
		section.gutter = Mm(0)

		doc.styles['Normal'].font.name = self.font if self.font else "Arial"

		if self.font_size:
			doc.styles['Normal'].font.size = Pt(self.font_size)

		for element in self.elements:
			element.doc_gen(doc)

		with open("out.xml", 'w+') as file:
			file.write(doc.element.xml)
			
		for e in doc.element.body:
			print(e)

		input()
		didPrint = False

		while True:
			try:
				doc.save(path)
				break
			except PermissionError:
				if not didPrint:
					PathTree.cls(False)
					print("Please close the document in order to save it...")
					didPrint = True
				
				time.sleep(0.5)

		PathTree.cls(False)
		print(f"Document saved to \"{path}\"")

		doc.save(path)

	@classmethod
	def wizard(cls, assign_context: bool) -> Optional['Document']:
		PathTree.cls()

		doc = cls()
		if assign_context: Context._instance.assign_document(doc)
		doc.edit()

		return doc
	
	def edit(self):
		ElementWizard.wizard(self.elements, "Doc", ignored_types_to_add=['Test'], callback=self.serialize_to_disk)

	@property
	def doc_name(self) -> str:
		return Path(self.context.project_path).name

	def __str__(self) -> str:
		return f"Document ({self.doc_name})"
	
	def __repr__(self) -> str:
		string = f"Document([{', '.join([repr(e) for e in self.elements])}]"

		if self.font: string += ", font=" + repr(self.font)
		if self.font_size: string += ", font_size=" + repr(self.font_size)

		return string + ")"